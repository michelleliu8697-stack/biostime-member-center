# -*- coding: utf-8 -*-
import docx
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

doc = docx.Document(r'D:/michelle.liu.luci.喵/品牌/合生元/合生元京东会员中心策划方案0413(1)(5).docx')
for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

# Also read tables
for table in doc.tables:
    for row in table.rows:
        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
        if row_text:
            print(' | '.join(row_text))
